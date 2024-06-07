import docx
from PIL import Image

# List of file paths
file_paths = [
    "/mnt/data/HOMO-27.bmp",
    "/mnt/data/HOMO-28.bmp",
    "/mnt/data/HOMO-29.bmp",
    "/mnt/data/HOMO-30.bmp"
]

# Create a new Word document
doc = docx.Document()

# Add a table with 2 columns
table = doc.add_table(rows=1, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'File Name'
hdr_cells[1].text = 'Cropped Image'

for file_path in file_paths:
    # Load the image
    image = Image.open(file_path)
    
    # Convert image to numpy array
    image_array = np.array(image)

    # Find the bounding box of the non-white areas
    non_white_pixels = np.where(np.all(image_array != [255, 255, 255], axis=-1))
    top_left = np.min(non_white_pixels, axis=1)
    bottom_right = np.max(non_white_pixels, axis=1)

    # Crop the image to the bounding box
    cropped_image_array = image_array[top_left[0]:bottom_right[0]+1, top_left[1]:bottom_right[1]+1]

    # Create a new image from the cropped array
    cropped_image = Image.fromarray(cropped_image_array)
    
    # Save the cropped image to a temporary file
    cropped_image_path = file_path.replace('.bmp', '_cropped.png')
    cropped_image.save(cropped_image_path)

    # Get the file name without extension
    file_name = file_path.split('/')[-1].replace('.bmp', '')

    # Add a row to the table
    row_cells = table.add_row().cells
    row_cells[0].text = file_name
    row_cells[1].add_paragraph().add_run().add_picture(cropped_image_path)

# Save the document
doc_output_path = "/mnt/data/cropped_images_table.docx"
doc.save(doc_output_path)

doc_output_path
